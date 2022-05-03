#!/usr/bin/env python
import ast
import csv
import json
import os
import re
from pathlib import Path
from textwrap import dedent
from zipfile import ZipFile

import click
import pandas as pd
from docx import Document

basedir = Path(__file__).resolve().parent
sourcedir = basedir / 'source'
mappingdir = basedir / 'output' / 'mapping'
eformsdir = mappingdir / 'eForms'

# We try to use pandas' N/A logic. However, in some cases we use `keep_default_na` to avoid
# "ValueError: Cannot mask with non-boolean array containing NA / NaN values".


def excel_files():
    with ZipFile(sourcedir / 'Task 5_Support_Standard Forms-eForms mappings_v.3.zip') as zipfile:
        for name in zipfile.namelist():
            with zipfile.open(name) as fileobj:
                with pd.ExcelFile(fileobj) as xlsx:
                    yield name, xlsx


def annex_rows():
    with pd.ExcelFile(sourcedir / 'CELEX_32019R1780_EN_ANNEX_TABLE2_Extended.xlsx') as xlsx:
        # A warning is issued, because the Excel file has an unsupported extension.
        df = pd.read_excel(xlsx, 'Annex')

        # Use the notice number columns.
        df.rename(columns={df.columns[i]: str(df.iloc[0][i]) for i in range(6, 51)}, errors='raise', inplace=True)

        for i, row in df.iterrows():
            if pd.notna(row['ID']):
                yield row


def write_guidance_pair(basename, df, **kwargs):
    df.to_csv(eformsdir / f'{basename}.csv', **kwargs)
    if not df['eforms_xpath'].apply(isinstance, args=[list]).any():
        df['eforms_xpath'] = df['eforms_xpath'].str.split(';')
    df.to_json(eformsdir / f'{basename}.json', orient='records', indent=2)


# https://github.com/pallets/click/issues/486
@click.group(context_settings={'max_content_width': 150})
def cli():
    pass


@cli.command()
@click.argument('sheet')
def find(sheet):
    """
    Find the workbook containing the SHEET.
    """
    for name, xlsx in excel_files():
        if sheet in xlsx.sheet_names:
            click.echo(f'First occurrence: {name}')
            break
    else:
        click.secho(f"Sheet '{sheet}' not found", fg='red')


@cli.command()
def extract_xpath_mapping():
    """
    Extract a mapping between Business Terms and eForms XPaths.

    \b
    Create or update output/mapping/eForms/bt-xpath-mapping.csv from source/XPATHs provisional release v. 1.0.docx
    """

    def text(row):
        # Newlines occur in all columns except the first, e.g.:
        # /*/cac:ContractingParty/cac:ContractingRepresentationType/cbc:RepresentationTypeCode
        # /*/cac:ProcurementProjectLot/cac:TenderingTerms/cac:TenderRecipientParty/cbc:EndpointID
        # Leading or trailing whitespace occur in the first and third columns, e.g.:
        # /*/cac:ProcurementProjectLot/cac:ProcurementProject/cbc:EstimatedOverallContractQuantity
        # /*/ext:UBLExtensions/ext:UBLExtension/ext:ExtensionContent/efext:EformsExtension/efac:Change/efbc:C...
        cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
        # XXX: Correct a typo.
        cells[0] = cells[0].replace('[@n listName=', '[@listName=')
        return cells

    docx = Document(sourcedir / 'XPATHs provisional release v. 1.0.docx')

    length = len(docx.tables)
    assert length == 1, f'unexpected table, found {length}'

    columns = text(docx.tables[0].rows[0])
    assert columns == ['XPATH', 'BT ID', 'BT Name', 'Additional information'], f'unexpected headers, got {columns}'

    data = []
    for row in docx.tables[0].rows[1:]:
        cells = text(row)
        xpath = cells[0]
        # Skip subheadings.
        if not xpath.startswith('/'):
            continue
        # XXX: Correct a typo (n-dash instead of empty).
        if cells[1] == '–':
            cells[1] = ''
        data.append(cells)

    pd.DataFrame(data, columns=columns).to_csv(eformsdir / 'bt-xpath-mapping.csv', index=False)


@cli.command()
def extract_indices_mapping():
    """
    Extract a mapping between Business Terms and form indices.

    \b
    Create or update output/mapping/eForms/bt-indices-mapping.csv from
    source/Task 5_Support_Standard Forms-eForms mappings_v.3.zip
    """
    ignore_sheets = {
        'Annex table 2',
        'eN40 vs F20',
        'Legend',
        'Legend Annex table 2',
        'S.F. vs eForms mapping list ',
    }
    ignore_sheet_regex = [re.compile(pattern) for pattern in (
        r'^F\d\d vs eN\d\d( \(2\))?$',
        r'^SF\d\d vs eForm ?\d\d(,\d\d)*$'
    )]
    keep_sheet_regex = re.compile(r'^(?:eForm|eN) ?(\d\d?(?:,\d\d)*) vs S?F(\d\d) ?$')

    replace_newlines = re.compile(r'\n(?=\(|(2009|Title III|and|requirements|subcontractor|will)\b)')
    explode = ['Level', 'Element']

    if 'DEBUG' in os.environ:
        (sourcedir / 'xlsx').mkdir(exist_ok=True)

    dfs = []
    for name, xlsx in excel_files():
        for sheet in xlsx.sheet_names:
            if sheet in ignore_sheets or any(regex.search(sheet) for regex in ignore_sheet_regex):
                continue

            match = keep_sheet_regex.search(sheet)
            if not match:
                raise click.ClickException(f"The sheet name {sheet!r} doesn't match a known pattern.")

            eforms_notice_number = match.group(1)
            sf_notice_number = match.group(2).lstrip('0')

            # Read the Excel file. The first row is a title for the table.
            df = pd.read_excel(xlsx, sheet, skiprows=[0], na_values='---')

            # Avoid empty or duplicate headings.
            df.rename(columns={
                df.columns[0]: 'Empty',  # was "" ("Unnamed: 0" after `read_excel`)
                df.columns[1]: 'Indent level',  # was "Level" or "" ("Unnamed: 1" after `read_excel`)
                df.columns[8]: 'Level',  # "Level.1" after `read_excel` if column 1 was "Level"
            }, errors='raise', inplace=True)

            # Remove rows with an empty "Level", for which we have no mapping information.
            df = df[df['Level'].notna()]

            # Remove the first column if it is empty.
            if df['Empty'].isna().all():
                df.drop(columns='Empty', inplace=True)
            else:
                raise click.ClickException("The first column was expected to be empty.")

            # Add notice number columns, using '1' instead of '01' to ease joins.
            df['eformsNotice'] = [[number.lstrip('0') for number in eforms_notice_number.split(',')] for i in df.index]
            df['sfNotice'] = sf_notice_number

            # Trim whitespace.
            df['Name'] = df['Name'].str.strip()

            # Make values easier to work with (must occur after `isna` above).
            df.fillna('', inplace=True)

            if 'DEBUG' in os.environ:
                df.to_csv(sourcedir / 'xlsx' / f'eForm{eforms_notice_number} - SF{sf_notice_number}.csv', index=False)

            # `explode` requires lists with the same number of elements, but an "Element" is not repeated if it is
            # the same for all "Level". Extra newlines also complicate things.
            for _, row in df.iterrows():
                location = f'eForm{eforms_notice_number} - SF{sf_notice_number}: {row["ID"]}: '

                # Remove spurious newlines in "Element" values.
                row['Element'] = row['Element'].replace('\n\n', '\n')
                if replace_newlines.search(row['Element']):
                    # Display the newlines that were removed, for user to review.
                    highlight = replace_newlines.sub(lambda m: click.style(m.group(0), blink=True), row['Element'])
                    # Replace outside f-string ("SyntaxError: f-string expression part cannot include a backslash").
                    highlight = highlight.replace('\n', '\\n')
                    click.echo(f'{location}removed \\n: {highlight}')
                    row['Element'] = replace_newlines.sub(' ', row['Element'])

                # XXX: Hardcode corrections or cases requiring human interpretation.
                if (
                    eforms_notice_number == '15'
                    and sf_notice_number == '7'
                    and row['ID'] == 'BT-18'
                    and row['Level'] == 'I.3.4.1.1'
                    and len(row['Element'].split('\n')) == 2
                ):
                    row['Level'] = 'I.3.4.1.1\nI.3.4.1.2'
                elif (
                    eforms_notice_number == '18'
                    and sf_notice_number == '17'
                    and row['ID'] == 'BT-750'
                    and row['Level'] == 'III.2.1.1.1\nIII.2.2.2\nIII.2.2.3\nIII.2.3.1.1\nIII.2.3.1.2'
                    and len(row['Element'].split('\n')) == 2
                ):
                    row['Element'] = dedent("""\
                    Information and formalities necessary for evaluating if the requirements are met:
                    Information and formalities necessary for evaluating if the requirements are met:
                    Minimum level(s) of standards possibly required: (if applicable)
                    Information and formalities necessary for evaluating if the requirements are met:
                    Minimum level(s) of standards possibly required: (if applicable)
                    """)

                if '\n' not in row['Level']:
                    # Warn about remaining newlines (BT-531 twice).
                    if '\n' in row['Element']:
                        click.secho(f'{location}unexpected \\n: {repr(row["Element"])}', fg='yellow')
                    continue

                # Split values, after removing leading and trailing newlines (occurs in "Level" values).
                for column in explode:
                    row[column] = row[column].strip("\n").split("\n")

                # Repeat "Element" as many times as there are "Level".
                length = len(row['Level'])
                if length > len(row['Element']):
                    row['Element'] *= length
                if length != len(row['Element']):
                    click.secho(f'{location}size differs: {row["Level"]} {row["Element"]}', fg='red')

                # Remove rows with a B... "Level" after split, for which we have no mapping information.
                for i, (level, element) in enumerate(zip(row['Level'], row['Element'])):
                    if level.startswith('B'):  # one row is "B 4.2" instead of "B.4.2"
                        del row['Level'][i]
                        del row['Element'][i]

            try:
                df = df.explode(explode)
            except ValueError as e:
                raise click.ClickException(f'{sheet}: {e}')

            df = df.explode('eformsNotice')

            dfs.append(df)

    pd.concat(dfs, ignore_index=True).to_csv(eformsdir / 'bt-indices-mapping.csv', index=False)


@cli.command()
def extract_hierarchy():
    """
    Extract the hierarchy of Business Groups and Business Terms.

    \b
    Create or update output/mapping/eForms/bt-bg-hierarchy.csv from
    source/CELEX_32019R1780_EN_ANNEX_TABLE2_Extended.xlsx
    """
    data = []
    line = []
    previous_level = 0

    for row in annex_rows():
        identifier = row['ID']
        current_level = len(row['Level'])

        # Adjust the size of this line of the "tree", then update the head.
        if current_level > previous_level:
            line.append(None)
        elif current_level < previous_level:
            line = line[:current_level]
        line[-1] = identifier

        data.append([identifier, *line[:-1]])
        previous_level = current_level

    pd.DataFrame(data, columns=['BT', 'BG_lvl1', 'BG_lvl2', 'BG_lvl3']).to_csv(
        eformsdir / 'bt-bg-hierarchy.csv', index=False
    )


@cli.command()
def update_ted_xml_indices():
    pass


@cli.command()
def extract_2015_guidance():
    """
    Concatenate guidance for the 2015 regulation.

    \b
    Create or update output/mapping/eForms/2015-guidance.csv:

    - Concatenate the CSV files for the 2015 regulation.
    - Merge the ted-xml-indices.csv file, which replaces the "index" column.
    """
    df_indices = pd.read_csv(eformsdir / 'ted-xml-indices.csv')

    dfs = []
    for path in mappingdir.glob('*.csv'):
        # The other columns are "index" and "comment".
        df = pd.read_csv(path, usecols=['xpath', 'label-key', 'guidance'])
        # Add the "index" column from the other file.
        df = df.merge(df_indices, how='left', on='xpath')
        # Add a "file" column for the source of the row.
        df['file'] = path.name
        dfs.append(df)

    # ignore_index is required, as each data frame repeats indices. Re-order the columns.
    pd.concat(dfs, ignore_index=True).to_csv(
        eformsdir / '2015-guidance.csv', columns=['xpath', 'label-key', 'index', 'guidance', 'file'], index=False
    )


@cli.command()
def prepopulate():
    """
    Prepopulate the guidance for the 2019 regulation, based on that for the 2015 regulation.

    \b
    Create or update output/mapping/eForms/2019-guidance-imported.json and .csv.
    """

    def add(data, current_row):
        current_row['XPATH'] = ';'.join(sorted(current_row['XPATH']))  # for briefer diff
        data.append(current_row)

    # Start with the eForms file that contains indices used by the 2015 guidance.
    # Without `dtype`, pandas writes the integers as floats (i.e. with decimals).
    df = pd.read_csv(eformsdir / 'bt-indices-mapping.csv', dtype={'eformsNotice': str, 'sfNotice': str})

    # Merge in the eForms XPaths. Sort by "BT ID" to simplify the for-loop's logic below.
    df_xpath = pd.read_csv(eformsdir / 'bt-xpath-mapping.csv').sort_values('BT ID')

    # 8 rows contain duplicate XPaths. Combine these with the original occurrences.
    data = []

    current_row = {'BT ID': None, 'XPATH': []}

    for _, row in df_xpath.iterrows():
        if row['BT ID'] != current_row['BT ID']:
            if current_row['XPATH']:
                add(data, current_row)
            row['XPATH'] = [row['XPATH']]
            current_row = row
        else:
            current_row['XPATH'].append(row['XPATH'])

    add(data, current_row)

    df = df.merge(pd.DataFrame(data, columns=df_xpath.columns), how='left', left_on='ID', right_on='BT ID')

    # TODO: This sort changes the output!
    df.sort_values(['ID', 'eformsNotice', 'sfNotice', 'Level', 'XPATH'], inplace=True)

    # Merge in hierarchy information and 2015 guidance.
    # TODO: Remove keep_default_na=False once comparison finished (changes sort order).
    df = df.merge(pd.read_csv(eformsdir / 'bt-bg-hierarchy.csv', keep_default_na=False), how='left', left_on='ID', right_on='BT')
    df = df.merge(pd.read_csv(eformsdir / '2015-guidance.csv'), how='left', left_on='Level', right_on='index')

    # TODO: Double-check this.
    df.drop_duplicates(['ID', 'eformsNotice'], inplace=True)

    # Avoid "ValueError: DataFrame columns must be unique for orient='records'."
    df.drop(columns='BT', inplace=True)  # same as "ID"

    df.loc[df['guidance'].notna(), 'status'] = 'imported_from_sf'
    df['comments'] = ''

    df.sort_values(['eformsNotice', 'BG_lvl1', 'BG_lvl2', 'BG_lvl3', 'ID'], inplace=True)

    df.rename(columns={
        'ID': 'BT',
        'Level': 'sfLevel',
        'XPATH': 'eforms_xpath',
    }, errors='raise', inplace=True)

    # Re-order columns. This effectively drops 13 columns:
    #
    # bt-indices-mapping.csv
    # - Indent level: superseded by bt-bg-mapping.csv
    # - Data type: re-added by sync-with-annex
    # - Repeatable
    # - Description: re-added by sync-with-annex
    # - Legal Status: re-added by sync-with-annex
    # - Element
    #
    # bt-xpath-mapping.csv
    # - BT ID: merge column
    # - BT Name: semantically the same as "Name"
    # - Additional information
    #
    # 2015-guidance.csv (only "guidance" needed)
    # - xpath
    # - label-key
    # - index: merge column
    # - file
    df = df[[
        'Name', 'eformsNotice', 'sfNotice', 'eforms_xpath', 'BT', 'sfLevel', 'BG_lvl1', 'BG_lvl2', 'BG_lvl3',
        'guidance', 'status', 'comments'
    ]]  # 12 columns

    write_guidance_pair('2019-guidance-imported', df, index=False)


@cli.command()
def update_with_annex():
    """
    Update the guidance with details from the 2019 regulation's annex.

    eforms-guidance.csv is read to write both eforms-guidance.csv and eforms-guidance.json.
    """
    df = pd.read_csv(
        eformsdir / 'eforms-guidance.csv',
        dtype={'eformsNotice': str, 'sfNotice': str},
        converters={'eforms_xpath': ast.literal_eval}  # TODO: Change the CSV to use ";"
    )

    df_annex = pd.DataFrame(annex_rows()).set_index('ID', verify_integrity=True)

    for label, row in df.iterrows():
        term = row['BT']
        status = df_annex.at[term, row['eformsNotice']]
        if status in ('CM', 'EM'):  # if "conditions" met or if it "exists"
            status = 'M'
        df.at[label, 'legal_status'] = status
        df.at[label, 'bt_description'] = df_annex.at[term, 'Description']
        df.at[label, 'bt_datatype'] = df_annex.at[term, 'Data type']

    for column in ('legal_status', 'guidance', 'status', 'comments'):  # TODO: Change to use null
        df[column].fillna('', inplace=True)

    df = df[[
        'id', 'Name', 'eformsNotice', 'sfNotice', 'legal_status', 'eforms_xpath', 'BT', 'bt_description',
        'bt_datatype', 'sfLevel', 'guidance', 'status', 'comments',
    ]]

    write_guidance_pair('eforms-guidance', df, index_label='id')


@cli.command()
def statistics():
    """
    Print statistics on the progress of the guidance for the 2019 regulation.
    """
    df = pd.read_csv(eformsdir / 'eforms-guidance.csv', keep_default_na=False)

    df_terms = df.drop_duplicates(subset='BT')
    total_terms = df_terms.shape[0]
    done_terms = df_terms[df_terms['status'].str.startswith('done')].shape[0]

    total = df.shape[0]
    imported = df[df['status'] == 'imported from standard forms'].shape[0]
    done = df[df['status'].str.startswith('done')].shape[0]
    ready = imported + done
    no_issue_no_guidance = df[(df['status'] == '') & (df['guidance'] == '')].shape[0]

    df_issue = df[df['status'].str.startswith('issue')]
    issue = df_issue.shape[0]
    issue_no_guidance = df_issue[df_issue['guidance'] == ''].shape[0]

    df_mandatory = df[df['legal_status'] == 'M']
    df_optional = df[df['legal_status'] == 'O']
    total_mandatory = df_mandatory.shape[0]
    total_optional = df_optional.shape[0]
    done_mandatory = df_mandatory[df_mandatory['status'].str.startswith('done')].shape[0]
    done_optional = df_optional[df_optional['status'].str.startswith('done')].shape[0]
    issue_mandatory = df_mandatory[df_mandatory['status'].str.startswith('issue')].shape[0]
    issue_optional = df_optional[df_optional['status'].str.startswith('issue')].shape[0]

    click.echo(dedent(f"""\
    - BTs ready for review: {done_terms}/{total_terms} ({done_terms / total_terms:.1%})
    - Rows ready for review: {ready}/{total} ({ready / total:.1%})
        - Imported from 2015 guidance: {imported} ({imported / total:.1%})
        - Added or edited after import: {done} ({done / total:.1%})
        - Per legal status:
            - Mandatory: {done_mandatory} ({done_mandatory / total_mandatory:.1%} of all M), {issue_mandatory} with open issues ({issue_mandatory / total_mandatory:.1%} of all M)
            - Optional: {done_optional} ({done_optional / total_optional:.1%} of all O), {issue_optional} with open issues ({issue_optional / total_optional:.1%} of all O)
    - Rows with [open issues](https://github.com/open-contracting/european-union-support/labels/eforms): {issue} ({issue / total:.1%}), {issue_no_guidance} without guidance
    - Rows without issues and without guidance: {no_issue_no_guidance} ({no_issue_no_guidance / total:.1%})
    """))  # noqa: E501


@cli.command()
@click.argument('file', type=click.File())
def fields_without_extensions(file):
    """
    Print undefined fields in the guidance for the 2015 regulation.
    """
    subjects = {
        # Unambiguous
        'award': 'awards',
        'contract': 'contracts',
        'lot': 'tender/lots',
        'party': 'parties',
        'release': '',
        'statistic': 'bids/statistics',
        'charge': 'contracts/implementation/charges',

        # Ambiguous
        'amendment': {
            'CHANGES': 'tender/amendments',
            'MODIFICATIONS_CONTRACT': 'contracts/amendments',
        },
        'classification': {
            'CONTRACTING_BODY': 'parties/details/classifications',
            'PROCEDURE': 'tender/procurementMethodRationaleClassifications',
        },
        'criterion': {
            'LEFTI': 'tender/selectionCriteria/criteria',
            'OBJECT_CONTRACT': 'tender/lots/awardCriteria/criteria',
        },
        'item': {
            'MODIFICATIONS_CONTRACT': 'contracts/items',
            'OBJECT_CONTRACT': 'tender/items',
        },
        'object': {
            '/OBJECT_CONTRACT/OBJECT_DESCR/EU_PROGR_RELATED': 'planning/budget/finance',
        },
    }

    unknowns = {
        # Unambiguous
        '.additionalContactPoints': 'parties',
        '.awardCriteria': 'tender',
        '.awardID': 'contracts',
        '.countryCode': 'parties/address',
        '.details.classifications': 'parties',
        '.documentType': 'tender/documents',
        '.estimatedValue.amount': 'contracts/implementation/charges',
        '.financingParty.id': 'planning/budget/finance',
        '.financingParty.name': 'planning/budget/finance',
        '.identifier.id': 'parties',
        '.identifier.legalName': 'parties',
        '.identifier.scheme': 'parties',
        '.measure': 'bids/statistics',  # metrics extension not used
        '.minimum': 'tender/selectionCriteria/criteria',
        '.paidBy': 'contracts/implementation/charges',
        '.roles': 'parties',
        '.secondStage.maximumCandidates': 'tender/lots',
        '.secondStage.minimumCandidates': 'tender/lots',
        '.subcontracting.maximumPercentage': 'awards',
        '.suppliers': 'awards',  # contract suppliers extension not used

        # Ambiguous
        '.additionalClassifications': {
            'MODIFICATIONS_CONTRACT': 'contracts/items',
        },
        '.description': {
            # Root
            'LEFTI': 'tender/selectionCriteria/criteria',
            'PROCEDURE': 'tender/procurementMethodRationaleClassifications',
            # XPath
            '/CONTRACTING_BODY/CA_ACTIVITY': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_ACTIVITY/@VALUE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_ACTIVITY_OTHER': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_TYPE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_TYPE/@VALUE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_TYPE_OTHER': 'parties/details/classifications',
            '/CONTRACTING_BODY/CE_ACTIVITY': 'parties/details/classifications',
            '/CONTRACTING_BODY/CE_ACTIVITY/@VALUE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CE_ACTIVITY_OTHER': 'parties/details/classifications',
            '/OBJECT_CONTRACT/CATEGORY': 'tender/additionalClassifications',
            '/OBJECT_CONTRACT/CATEGORY/@CTYPE': 'tender/additionalClassifications',
            '/OBJECT_CONTRACT/OBJECT_DESCR/EU_PROGR_RELATED': 'planning/budget/finance',
        },
        '.id': {
            # Root
            'CHANGES': 'tender/amendments',
            'LEFTI': 'tender/documents',
            'PROCEDURE': 'tender/procurementMethodRationaleClassifications',
            # XPath
            '/AWARD_CONTRACT': 'awards',
            '/AWARD_CONTRACT/AWARDED_CONTRACT': 'contracts',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/CONTRACTORS/CONTRACTOR/ADDRESS_CONTRACTOR': 'awards/suppliers',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/CONTRACTORS/CONTRACTOR/ADDRESS_PARTY': 'parties/shareholders',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/TENDERS/NB_TENDERS_RECEIVED': 'bids/statistics',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/TENDERS/NB_TENDERS_RECEIVED_EMEANS': 'bids/statistics',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/TENDERS/NB_TENDERS_RECEIVED_NON_EU': 'bids/statistics',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/TENDERS/NB_TENDERS_RECEIVED_OTHER_EU': 'bids/statistics',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/TENDERS/NB_TENDERS_RECEIVED_SME': 'bids/statistics',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/VAL_PRICE_PAYMENT': 'contracts/implementation/charges',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/VAL_REVENUE': 'contracts/implementation/charges',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/VALUES/VAL_RANGE_TOTAL/HIGH': 'bids/statistics',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/VALUES/VAL_RANGE_TOTAL/LOW': 'bids/statistics',
            '/CONTRACTING_BODY/ADDRESS_CONTRACTING_BODY': 'parties',
            '/CONTRACTING_BODY/DOCUMENT_RESTRICTED': 'tender/participationFees',
            '/CONTRACTING_BODY/CA_ACTIVITY': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_ACTIVITY/@VALUE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_ACTIVITY_OTHER': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_TYPE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_TYPE/@VALUE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_TYPE_OTHER': 'parties/details/classifications',
            '/CONTRACTING_BODY/CE_ACTIVITY': 'parties/details/classifications',
            '/CONTRACTING_BODY/CE_ACTIVITY/@VALUE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CE_ACTIVITY_OTHER': 'parties/details/classifications',
            '/MODIFICATIONS_CONTRACT/DESCRIPTION_PROCUREMENT/CONTRACTORS/CONTRACTOR/ADDRESS_CONTRACTOR': 'awards/suppliers',  # noqa: E501
            '/MODIFICATIONS_CONTRACT/DESCRIPTION_PROCUREMENT/CPV_ADDITIONAL/CPV_CODE': 'contracts/items/additionalClassifications',  # noqa: E501
            '/MODIFICATIONS_CONTRACT/DESCRIPTION_PROCUREMENT/CPV_ADDITIONAL/CPV_SUPPLEMENTARY_CODE': 'contracts/items/additionalClassifications',  # noqa: E501
            '/MODIFICATIONS_CONTRACT/DESCRIPTION_PROCUREMENT/CPV_MAIN/CPV_SUPPLEMENTARY_CODE': 'contracts/items/additionalClassifications',  # noqa: E501
            '/MODIFICATIONS_CONTRACT/INFO_MODIFICATIONS': 'contracts/amendments',
            '/OBJECT_CONTRACT/CATEGORY': 'tender/additionalClassifications',
            '/OBJECT_CONTRACT/CATEGORY/@CTYPE': 'tender/additionalClassifications',
            '/OBJECT_CONTRACT/CPV_MAIN/CPV_SUPPLEMENTARY_CODE': 'tender/items/additionalClassifications',
            '/OBJECT_CONTRACT/OBJECT_DESCR/CPV_ADDITIONAL/CPV_CODE': 'tender/items/additionalClassifications',
            '/OBJECT_CONTRACT/OBJECT_DESCR/CPV_ADDITIONAL/CPV_SUPPLEMENTARY_CODE': 'tender/items/additionalClassifications',  # noqa: E501
            '/OBJECT_CONTRACT/OBJECT_DESCR/EU_PROGR_RELATED': 'planning/budget/finance',
            '/OBJECT_CONTRACT/VAL_RANGE_TOTAL/HIGH': 'bids/statistics',
            '/OBJECT_CONTRACT/VAL_RANGE_TOTAL/LOW': 'bids/statistics',
            '/RESULTS/AWARDED_PRIZE': 'contracts',
            '/RESULTS/AWARDED_PRIZE/PARTICIPANTS/NB_PARTICIPANTS': 'bids/statistics',
            '/RESULTS/AWARDED_PRIZE/PARTICIPANTS/NB_PARTICIPANTS_OTHER_EU': 'bids/statistics',
            '/RESULTS/AWARDED_PRIZE/PARTICIPANTS/NB_PARTICIPANTS_SME': 'bids/statistics',
            '/RESULTS/AWARDED_PRIZE/WINNERS/WINNER/ADDRESS_WINNER': 'awards/suppliers',
        },
        '.items': {
            'MODIFICATIONS_CONTRACT': 'contracts',
        },
        '.name': {
            # Root
            'AWARD_CONTRACT': 'awards/suppliers',
            'MODIFICATIONS_CONTRACT': 'awards/suppliers',
            'OBJECT_CONTRACT': 'parties',
            'RESULTS': 'awards/suppliers',
            # XPath
            '/CONTRACTING_BODY/ADDRESS_CONTRACTING_BODY': 'buyer',
            '/PROCEDURE/PARTICIPANT_NAME': 'parties',
            '/PROCEDURE/MEMBER_NAME': 'tender/designContest/juryMembers',
        },
        '.newValue': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.newValue.classifications': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.newValue.date': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.newValue.text': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.oldValue': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.oldValue.classifications': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.oldValue.date': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.oldValue.text': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.region': {
            '/OBJECT_CONTRACT/OBJECT_DESCR/NUTS': 'tender/items/deliveryAddresses',
            '/MODIFICATIONS_CONTRACT/DESCRIPTION_PROCUREMENT/NUTS': 'contracts/items/deliveryAddresses',
        },
        '.relatedLot': {
            'AWARD_CONTRACT': 'bids/statistics',
            'CHANGES': 'tender/amendments/unstructuredChanges',
            'RESULTS': 'bids/statistics',
        },
        '.relatedLots': {
            'AWARD_CONTRACT': 'awards',
            'OBJECT_CONTRACT': 'planning/budget/finance',
        },
        '.scheme': {
            'MODIFICATIONS_CONTRACT': 'contracts/items/additionalClassifications',
            # XPath
            '/CHANGES/CHANGE/NEW_VALUE/CPV_ADDITIONAL/CPV_CODE': 'tender/amendments/unstructuredChanges/newValue/classifications',  # noqa: E501
            '/CHANGES/CHANGE/NEW_VALUE/CPV_ADDITIONAL/CPV_SUPPLEMENTARY_CODE': 'tender/amendments/unstructuredChanges/newValue/classifications',  # noqa: E501
            '/CHANGES/CHANGE/NEW_VALUE/CPV_MAIN/CPV_CODE': 'tender/amendments/unstructuredChanges/newValue/classifications',  # noqa: E501
            '/CHANGES/CHANGE/NEW_VALUE/CPV_MAIN/CPV_SUPPLEMENTARY_CODE': 'tender/amendments/unstructuredChanges/newValue/classifications',  # noqa: E501
            '/CHANGES/CHANGE/OLD_VALUE/CPV_ADDITIONAL/CPV_CODE': 'tender/amendments/unstructuredChanges/oldValue/classifications',  # noqa: E501
            '/CHANGES/CHANGE/OLD_VALUE/CPV_ADDITIONAL/CPV_SUPPLEMENTARY_CODE': 'tender/amendments/unstructuredChanges/oldValue/classifications',  # noqa: E501
            '/CHANGES/CHANGE/OLD_VALUE/CPV_MAIN/CPV_CODE': 'tender/amendments/unstructuredChanges/oldValue/classifications',  # noqa: E501
            '/CHANGES/CHANGE/OLD_VALUE/CPV_MAIN/CPV_SUPPLEMENTARY_CODE': 'tender/amendments/unstructuredChanges/oldValue/classifications',  # noqa: E501
            '/OBJECT_CONTRACT/CPV_MAIN/CPV_SUPPLEMENTARY_CODE': 'tender/items/additionalClassifications',
            '/OBJECT_CONTRACT/OBJECT_DESCR/CPV_ADDITIONAL/CPV_CODE': 'tender/items/additionalClassifications',
            '/OBJECT_CONTRACT/OBJECT_DESCR/CPV_ADDITIONAL/CPV_SUPPLEMENTARY_CODE': 'tender/items/additionalClassifications',  # noqa: E501
            '/OBJECT_CONTRACT/CATEGORY': 'tender/additionalClassifications',
        },
        '.shareholder.id': {
            'AWARD_CONTRACT': 'parties/shareholders',
        },
        '.shareholder.name': {
            'AWARD_CONTRACT': 'parties/shareholders',
        },
        '.status': {
            'AWARD_CONTRACT': 'contracts',
            'RESULTS': 'contracts',
        },
        '.title': {
            'AWARD_CONTRACT': 'contracts',
            'LEFTI': 'tender/targets',
        },
        '.type': {
            # Root
            'CONTRACTING_BODY': 'tender/participationFees',
            'LEFTI': 'tender.selectionCriteria.criteria',
            'OBJECT_CONTRACT': 'tender/lots/awardCriteria/criteria',
        },
        '.value': {
            'AWARD_CONTRACT': 'bids/statistics',
            'OBJECT_CONTRACT': 'bids/statistics',
            'RESULTS': 'bids/statistics',
        },
        '.where': {
            'CHANGES': 'tender/amendments/unstructuredChanges',
        },
        '.where.label': {
            'CHANGES': 'tender/amendments/unstructuredChanges',
        },
        '.where.section': {
            'CHANGES': 'tender/amendments/unstructuredChanges',
        },
    }

    unhandled = set()
    paths = set()

    def report(path, row):
        value = (path, row.get('xpath'))
        if value not in unhandled:
            print(f"unhandled: {path} ({row.get('xpath')}: {row['guidance']})", err=True)
        unhandled.add(value)

    for path in (mappingdir, mappingdir / 'shared'):
        for filename in path.glob('*.csv'):
            with filename.open() as f:
                reader = csv.DictReader(f)
                for row in reader:
                    if row.get('guidance'):
                        for match in re.finditer(r"(?:([a-z]+)'s )?\[?`([^`]+)`", row['guidance']):
                            path = match.group(2)
                            if path in ('true', 'false', 'value'):  # JSON boolean, exceptional case
                                continue
                            if re.search(r'^[A-Z][a-z][A-Za-z]+$', path):  # JSON Schema definition
                                continue
                            if re.search(r'^(/[A-Z_]+)+$', path):  # XPath
                                continue
                            if re.search(r'^[A-Z_]+$', path):  # XML element
                                continue

                            subject = match.group(1)

                            prefix = ''
                            if subject:
                                try:
                                    prefix = subjects[subject]
                                except KeyError as e:
                                    click.echo(f"KeyError: Add a {e} key to the `subjects` list")
                            elif path in unknowns:
                                try:
                                    prefix = unknowns[path]
                                except KeyError as e:
                                    click.echo(f"KeyError: Add a {e} key to the `unknowns` list")
                            elif path[0] == '.':
                                report(path, row)
                                continue

                            if isinstance(prefix, dict):
                                xpath = row.get('xpath', '/')
                                root = xpath.split('/', 2)[1]
                                if root in prefix:
                                    key = root
                                else:
                                    key = xpath
                                try:
                                    prefix = prefix[key]
                                except KeyError:
                                    report(path, row)
                                    continue

                            path = prefix + path
                            paths.add(path.replace('.', '/'))

    seen = [row['path'] for row in csv.DictReader(file)]
    for path in sorted(list(paths)):
        if path not in seen:
            click.echo(path)

    # Uncomment to print all the paths for a specific object.
    # for path in sorted(list(paths)):
    #     if '/items' in path:
    #         click.echo(path)


if __name__ == '__main__':
    cli()
